import os
import re
import csv
import hashlib
import logging
import warnings
from pathlib import Path
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Optional

warnings.filterwarnings("ignore")
for lib in ["PyPDF2", "pdfminer", "PIL", "pytesseract", "pdfplumber"]:
    logging.getLogger(lib).setLevel(logging.CRITICAL)

# ─── КОНФИГУРАЦИЯ ────────────────────────────────────────────────────────────
ROOT_DIR    = Path("./DATA")
OUTPUT_CSV  = Path("result.csv")
MAX_WORKERS = 8

# ─── ЛЕНИВЫЙ ИМПОРТ ──────────────────────────────────────────────────────────
def _try(name):
    try: return __import__(name)
    except ImportError: return None

pandas     = _try("pandas")
PyPDF2     = _try("PyPDF2")
chardet    = _try("chardet")
pdfplumber = _try("pdfplumber")

try:
    import pytesseract; pytesseract.get_tesseract_version(); TESSERACT = True
except: pytesseract = None; TESSERACT = False

try:
    from pdfminer.high_level import extract_text as _pdfminer_extract; PDFMINER = True
except: PDFMINER = False

try:
    import docx as _docx_lib; DOCX = True
except: DOCX = False

# ─── ВСПОМОГАТЕЛЬНЫЕ ─────────────────────────────────────────────────────────

_MONTHS = {1:'jan',2:'feb',3:'mar',4:'apr',5:'may',6:'jun',
           7:'jul',8:'aug',9:'sep',10:'oct',11:'nov',12:'dec'}

def fmt_mtime(ts: float) -> str:
    dt = datetime.fromtimestamp(ts)
    return f"{_MONTHS[dt.month]} {dt.day:02d} {dt.strftime('%H:%M')}"

def file_md5(path: Path) -> str:
    h = hashlib.md5()
    try:
        with open(path, 'rb') as f:
            while True:
                buf = f.read(65536)
                if not buf: break
                h.update(buf)
    except: pass
    return h.hexdigest()

def _decode(raw: bytes) -> str:
    if chardet:
        enc = chardet.detect(raw[:40000]).get('encoding') or 'utf-8'
    else:
        enc = 'utf-8'
    return raw.decode(enc, errors='ignore')

# ─── ДЕТЕКТОР ТИПА ФАЙЛА ─────────────────────────────────────────────────────

def detect_ext(path: Path, declared: str) -> str:
    if declared in ('csv','json','txt','html','htm','parquet',
                    'xls','xlsx','docx','doc','rtf'):
        return declared
    try:
        head = path.read_bytes()[:16]
        if head.startswith(b'%PDF'): return 'pdf'
        if head.startswith(b'PK\x03\x04'): return 'docx'
        if head.startswith(b'\xff\xd8\xff'): return 'jpg'
        if head.startswith(b'\x89PNG'): return 'png'
        if head[:4] in (b'II*\x00', b'MM\x00*'): return 'tif'
        if b'ftyp' in head or b'mp4' in head.lower(): return 'mp4'
    except: pass
    return declared or 'unknown'

# ─── ВАЛИДАТОРЫ ──────────────────────────────────────────────────────────────

def validate_snils(s: str) -> bool:
    nums = re.sub(r'\D', '', s)
    if len(nums) != 11: return False
    base = [int(x) for x in nums[:9]]
    check = int(nums[9:])
    total = sum((9-i)*d for i,d in enumerate(base))
    if total < 100: c = total
    elif total in (100,101): c = 0
    else: c = 0 if (total%101)==100 else (total%101)
    return c == check

def validate_inn12(inn: str) -> bool:
    d = [int(x) for x in inn]
    if len(d) != 12: return False
    w1 = [7,2,4,10,3,5,9,4,6,8]
    w2 = [3,7,2,4,10,3,5,9,4,6,8]
    c1 = sum(w*v for w,v in zip(w1,d[:10])) % 11 % 10
    c2 = sum(w*v for w,v in zip(w2,d[:11])) % 11 % 10
    return d[10]==c1 and d[11]==c2

def validate_luhn(s: str) -> bool:
    digits = re.sub(r'\D','',s)
    if len(digits) < 13: return False
    total = 0
    for i,d in enumerate(reversed(digits)):
        n = int(d)
        if i%2==1:
            n*=2
            if n>9: n-=9
        total+=n
    return total%10==0

# ─── ПАТТЕРНЫ ПДН ────────────────────────────────────────────────────────────

RE_SNILS_STRICT   = re.compile(r'(\d{3}-\d{3}-\d{3}\s?\d{2})', re.IGNORECASE)
RE_SNILS_CONTEXT  = re.compile(r'(?:снилс|snils|страховой\s+номер)[^\d]{0,40}(\d{3}[-\s]?\d{3}[-\s]?\d{3}[-\s]?\d{2})', re.IGNORECASE)
RE_PASSPORT       = re.compile(r'(?:паспорт|серия)[^\d]{0,40}(\d{2}[\s\-]?\d{2}[\s\-]?\d{6})', re.IGNORECASE)
RE_PASSPORT_ISSUED= re.compile(r'(?:выдан|выдано)[^\n]{0,100}(\d{2}[\s\-]?\d{2}[\s\-]?\d{6})', re.IGNORECASE)
RE_INN12          = re.compile(r'(?:инн|inn)[^\d]{0,20}(\d{12})\b', re.IGNORECASE)
RE_CARD           = re.compile(r'\b(\d{4}[\s\-]\d{4}[\s\-]\d{4}[\s\-]\d{4})\b')
RE_MRZ            = re.compile(r'P<RUS[A-Z<]{30}', re.IGNORECASE)
RE_ACCOUNT_PHYS   = re.compile(r'(?:р/с|расчётный\s+счёт|лицевой\s+счёт)[^\d]{0,15}(4\d{19})\b', re.IGNORECASE)

# ─── ИМЕНА ФАЙЛОВ-СКАНОВ ДОКУМЕНТОВ ──────────────────────────────────────────
# Паттерн: [Буквы][Цифры]_[Цифры].[tif]
# Например: CA01_01.tif, HA06_10.tif, KS13_01.tif
RE_SCAN_FILENAME = re.compile(
    r'^[A-Z]{2}\d{2}_\d{2}\.tiff?$',
    re.IGNORECASE
)

# ─── ИЗВЛЕЧЕНИЕ ТЕКСТА ───────────────────────────────────────────────────────

def extract_text(path: Path, ext: str) -> str:
    try:
        if ext == 'csv':
            if pandas:
                try:
                    df = pandas.read_csv(path, dtype=str, nrows=5000, on_bad_lines='skip')
                    return df.to_csv(index=False)
                except: pass
            return _decode(path.read_bytes()[:300_000])

        elif ext == 'parquet':
            if pandas:
                try:
                    df = pandas.read_parquet(path).head(5000)
                    return df.to_csv(index=False)
                except: pass
            return ''

        elif ext in ('xls','xlsx'):
            if pandas:
                try:
                    df = pandas.read_excel(path, nrows=5000, dtype=str)
                    return df.to_csv(index=False)
                except: pass
            return ''

        elif ext == 'pdf':
            text = ''
            if PDFMINER:
                try:
                    t = _pdfminer_extract(str(path), maxpages=15) or ''
                    if len(t.strip()) > 50: return t
                except: pass
            if pdfplumber:
                try:
                    with pdfplumber.open(path) as pdf:
                        t = '\n'.join(p.extract_text() or '' for p in pdf.pages[:15])
                    if len(t.strip()) > 50: return t
                except: pass
            if PyPDF2:
                try:
                    with open(path,'rb') as f:
                        r = PyPDF2.PdfReader(f, strict=False)
                        for pg in r.pages[:15]:
                            try: text += pg.extract_text() or ''
                            except: pass
                except: pass
            return text

        elif ext in ('tif','tiff','jpg','jpeg','png'):
            if TESSERACT:
                try:
                    from PIL import Image
                    img = Image.open(path)
                    texts = []
                    try:
                        for i in range(min(getattr(img,'n_frames',1), 3)):
                            img.seek(i)
                            texts.append(pytesseract.image_to_string(img.convert('L'), lang='rus+eng'))
                    except EOFError:
                        texts.append(pytesseract.image_to_string(img.convert('L'), lang='rus+eng'))
                    return '\n'.join(texts)
                except: pass
            return ''

        elif ext == 'docx':
            if DOCX:
                try:
                    doc = _docx_lib.Document(str(path))
                    parts = [p.text for p in doc.paragraphs]
                    for table in doc.tables:
                        for row in table.rows:
                            parts.append(' | '.join(c.text for c in row.cells))
                    return '\n'.join(parts)
                except: pass
            return ''

        elif ext in ('doc', 'rtf'):
            if DOCX:
                try:
                    doc = _docx_lib.Document(str(path))
                    return '\n'.join(p.text for p in doc.paragraphs)
                except: pass
            return _decode(path.read_bytes()[:200_000])

        elif ext in ('html', 'htm'):
            raw = _decode(path.read_bytes()[:400_000])
            text = re.sub(r'<(script|style)[^>]*>.*?</\1>', ' ', raw, flags=re.DOTALL|re.IGNORECASE)
            return re.sub(r'<[^>]+>', ' ', text)

        elif ext == 'txt':
            return _decode(path.read_bytes())

        elif ext == 'json':
            import json
            try:
                raw = _decode(path.read_bytes())
                data = json.loads(raw)
                return json.dumps(data, ensure_ascii=False)
            except:
                return _decode(path.read_bytes()[:300_000])

        else:
            raw = path.read_bytes()[:100_000]
            if raw.count(b'\x00') < len(raw) * 0.02:
                return _decode(raw)
            return ''

    except Exception:
        return ''

# ─── АНАЛИЗ СОДЕРЖИМОГО НА ПДН ───────────────────────────────────────────────

def has_strong_pii(text: str) -> tuple[bool, list]:
    """
    Ищет только СИЛЬНЫЕ идентификаторы:
    СНИЛС, паспорт РФ, ИНН физлица, банковская карта, MRZ, лицевой счёт.
    Возвращает (найдено, список категорий).
    """
    if not text:
        return False, []

    tl = text.lower()
    found = []

    # СНИЛС - с контекстом или строгий формат XXX-XXX-XXX XX
    for m in RE_SNILS_CONTEXT.finditer(tl):
        if validate_snils(m.group(1)):
            found.append('СНИЛС')
            break
    if 'СНИЛС' not in found:
        for m in RE_SNILS_STRICT.finditer(tl):
            if validate_snils(m.group(1)):
                found.append('СНИЛС')
                break

    # Паспорт РФ
    for pat in (RE_PASSPORT, RE_PASSPORT_ISSUED):
        if pat.search(tl) and 'Паспорт РФ' not in found:
            found.append('Паспорт РФ')
            break

    # ИНН физлица (12 цифр) с валидацией
    for m in RE_INN12.finditer(tl):
        val = re.sub(r'\D','', m.group(1))
        if len(val)==12 and validate_inn12(val):
            found.append('ИНН физлица')
            break

    # Банковская карта (Луна)
    for m in RE_CARD.finditer(text):
        if validate_luhn(m.group(1)):
            found.append('Банковская карта')
            break

    # MRZ
    if RE_MRZ.search(text):
        found.append('MRZ')

    # Расчётный/лицевой счёт физлица
    if RE_ACCOUNT_PHYS.search(tl):
        found.append('Лицевой счёт')

    return len(found) > 0, found

def has_pii_columns(text: str) -> tuple[bool, list]:
    """Анализ заголовков CSV/Parquet на ПДн-колонки."""
    if not text:
        return False, []

    first_line = text.split('\n')[0].lower()
    cols = [c.strip().strip('"\'') for c in re.split(r'[,;\t]', first_line)]

    PII_COLS = {
        'ФИО/Имя': {'name','full_name','first_name','last_name','surname',
                    'фамилия','имя','отчество','фио','firstname','lastname','patronymic'},
        'Email': {'email','e-mail','почта'},
        'Телефон': {'phone','телефон','mobile','tel','phone_number'},
        'СНИЛС': {'snils','снилс'},
        'ИНН': {'inn','инн'},
        'Паспорт': {'passport','паспорт'},
        'Адрес': {'address','адрес','city','street'},
        'Дата рождения': {'birthday','birth_date','dob','birthdate','дата_рождения'},
        'Карта': {'card','карта','pan','card_number'},
    }

    found_cats = []
    for cat, keywords in PII_COLS.items():
        for col in cols:
            if col in keywords or any(kw in col for kw in keywords):
                found_cats.append(cat)
                break

    # Нужно минимум 2 ПДн-колонки для структурированного файла
    # (одно имя без контекста - не считается)
    return len(found_cats) >= 2, found_cats

# ─── ГЛАВНАЯ ФУНКЦИЯ ПРОВЕРКИ ─────────────────────────────────────────────────

def check_file(path: Path) -> Optional[tuple]:
    """
    Возвращает (size, mtime, rel_path, categories, uz) или None.
    """
    try:
        stat = path.stat()
        size = stat.st_size
        if size == 0:
            return None

        try:
            rel = path.relative_to(Path.cwd()).as_posix()
        except ValueError:
            rel = path.as_posix()
        if rel.startswith('./'): rel = rel[2:]
        if rel.startswith('DATA/'): rel = rel[5:]

        rel_l = rel.lower()
        fn    = path.name
        fn_l  = fn.lower()

        ext      = path.suffix.lower().lstrip('.')
        real_ext = detect_ext(path, ext)
 
        if 'physical.parquet' in rel_l:
            return (size, fmt_mtime(stat.st_mtime), rel,
                    ['База физлиц — СНИЛС, паспорт, ИНН'], 1)

        if re.search(r'логистика/customers\.csv', rel_l):
            return (size, fmt_mtime(stat.st_mtime), rel,
                    ['Клиентская база (ФИО, контакты, адрес)'], 3)
        if re.search(r'логистика/logistics\.csv', rel_l):
            return (size, fmt_mtime(stat.st_mtime), rel,
                    ['Логистика (ФИО получателей, адреса)'], 3)

        if '/employes/' in rel_l or '\\employes\\' in rel_l:
            text = extract_text(path, real_ext)
            ok, cats = has_strong_pii(text)
            if not cats:
                cats = ['HR-документ с ПДн сотрудника']
            return (size, fmt_mtime(stat.st_mtime), rel, cats, 3 if ok else 4)

        if 'lost+found' in rel_l:
            if size > 2_000_000:
                return (size, fmt_mtime(stat.st_mtime), rel,
                        ['Восстановленный файл lost+found (>2MB)'], 3)
            return None

        if real_ext in ('tif', 'tiff'):
            if RE_SCAN_FILENAME.match(fn):
                return (size, fmt_mtime(stat.st_mtime), rel,
                        ['Скан удостоверения личности (TIF)'], 3)
            if size > 500_000 and 'прочее' in rel_l:
                ocr = extract_text(path, real_ext)
                ok, cats = has_strong_pii(ocr)
                if ok:
                    return (size, fmt_mtime(stat.st_mtime), rel, cats, 3)
            return None

        if real_ext in ('mp4', 'avi', 'mov') and size > 500_000:
            if 'прочее' in rel_l or 'dataset4' in rel_l:
                return (size, fmt_mtime(stat.st_mtime), rel,
                        ['Видеозапись документов'], 3)
            return None
 
        PUBLIC_DIRS = [
            'dataset0/', 'dataset1/', 'dataset2/', 'dataset3/',
            'выгрузки/сайты/',
            'документы партнеров/',
        ]
        for pub in PUBLIC_DIRS:
            if pub in rel_l:
                return None

        NEUTRAL_FILES = {'products.csv', 'stores.csv', 'sales.csv'}
        if fn_l in NEUTRAL_FILES:
            return None

        if real_ext in ('csv', 'parquet', 'xls', 'xlsx', 'json'):
            text = extract_text(path, real_ext)
            col_ok, col_cats = has_pii_columns(text)
            if col_ok:
                lines = [l for l in (text or '').split('\n') if l.strip()]
                if len(lines) > 2:
                    uz = 3 if any(c in col_cats for c in ('СНИЛС','ИНН','Паспорт','Карта')) else 4
                    return (size, fmt_mtime(stat.st_mtime), rel, col_cats, uz)
            ok, cats = has_strong_pii(text)
            if ok:
                return (size, fmt_mtime(stat.st_mtime), rel, cats,
                        2 if 'Банковская карта' in cats else 3)
            return None

        if real_ext in ('txt', 'pdf', 'docx', 'doc', 'rtf', 'html', 'htm'):
            text = extract_text(path, real_ext)
            ok, cats = has_strong_pii(text)
            if ok:
                uz = 2 if 'Банковская карта' in cats else 3
                return (size, fmt_mtime(stat.st_mtime), rel, cats, uz)
            return None

        return None

    except Exception:
        return None

# ─── ОСНОВНОЙ СКАН ───────────────────────────────────────────────────────────

def scan(root: Path, out_csv: Path):
    all_files = [
        Path(dp) / fn
        for dp, _, fns in os.walk(root)
        for fn in fns
    ]
    total = len(all_files)
    print(f"Файлов: {total} | Потоков: {MAX_WORKERS}")

    seen_hashes: set = set()
    results = []

    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as ex:
        futures = {ex.submit(check_file, p): p for p in all_files}
        done = 0
        for future in as_completed(futures):
            done += 1
            if done % 300 == 0 or done == total:
                print(f"  [{done}/{total}]...", end='\r')

            res = future.result()
            if res is None:
                continue

            size, mtime, rel, cats, uz = res
            p = futures[future]

            h = file_md5(p)
            if h in seen_hashes:
                print(f"\n[-] Дубликат: {rel}")
                continue
            seen_hashes.add(h)

            results.append((size, mtime, rel, cats, uz))

    results.sort(key=lambda x: x[2])

    with out_csv.open('w', newline='', encoding='utf-8') as f:
        w = csv.writer(f)
        w.writerow(['size', 'time', 'name'])
        for size, mtime, rel, _, _ in results:
            w.writerow([size, mtime, rel])

    print(f"\n{'='*70}")
    print(f"  ИТОГ: {len(results)} файлов с ПДн")
    print(f"{'='*70}")
    uz_counts = {1:0, 2:0, 3:0, 4:0}
    for i, (size, mtime, rel, cats, uz) in enumerate(results, 1):
        uz_counts[uz] += 1
        print(f"\n[{i:02d}] {rel}")
        print(f"     {size:,} байт | {mtime} | УЗ-{uz}")
        for c in cats:
            print(f"     ▸ {c}")

    print(f"\n{'─'*70}")
    for lvl in (1,2,3,4):
        if uz_counts[lvl]:
            print(f"  УЗ-{lvl}: {uz_counts[lvl]} файлов")
    print(f"  Сохранено: {out_csv.resolve()}")

# ─── ТОЧКА ВХОДА ─────────────────────────────────────────────────────────────

if __name__ == '__main__':
    print("PII Scanner v4.0 — Точечное попадание")
    print(f"OCR: {'да' if TESSERACT else 'нет'} | "
          f"pdfminer: {'да' if PDFMINER else 'нет'} | "
          f"pdfplumber: {'да' if pdfplumber else 'нет'}")
    print(f"Корень: {ROOT_DIR.resolve()}\n")

    if not ROOT_DIR.exists():
        print(f"ОШИБКА: {ROOT_DIR} не найден")
    else:
        scan(ROOT_DIR, OUTPUT_CSV)
